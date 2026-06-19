from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Style config --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# -- Helper functions --
def add_heading_custom(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x2A)  # forest green
    return h

def add_section_heading(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(8)
    run = h.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x2A)
    return h

def add_sub_heading(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x2A)
    return h

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_feature(name, desc):
    p = doc.add_paragraph()
    run_name = p.add_run(name + ': ')
    run_name.bold = True
    run_name.font.size = Pt(11)
    p.add_run(desc)

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for r in paragraph.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    return table

# ============================ TITLE ============================
title = doc.add_heading('DAFTAR FITUR ECOLIFE HUB', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x2A)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Sustainable Living Platform\n\n')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x54, 0x6E, 0x4C)
run2 = subtitle.add_run('Dokumentasi Fitur - Semua Hak Cipta')
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(0x7A, 0x8A, 0x7A)

doc.add_page_break()

# ============================ TABLE OF CONTENTS ============================
add_heading_custom('DAFTAR ISI', level=1)
toc_items = [
    'A. Fitur untuk Guest (Belum Login)',
    'B. Fitur untuk User (Sudah Login)',
    '   B1. Dashboard',
    '   B2. Nutrisi / Food Tracking',
    '   B3. Aktivitas / Olahraga',
    '   B4. Daily Report',
    '   B5. History',
    '   B6. Artikel Edukasi',
    '   B7. Forum Diskusi',
    '   B8. Quiz SDG',
    '   B9. Achievements',
    '   B10. SDG Detail',
    '   B11. Profil',
    '   B12. Body Data',
    'C. Fitur Admin',
    '   C1. Admin Dashboard',
    '   C2. Manage Articles',
    '   C3. Manage Users',
    '   C4. Manage Comments',
    '   C5. Manage Discussions',
    '   C6. Manage Quiz Questions',
    '   C7. Data Overview',
    'D. Fitur Teknis / Lintas Fitur',
    '   D1. Multi Bahasa',
    '   D2. Achievement System',
    '   D3. AI Integration',
    '   D4. Tiga-tier Food Search',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ============================ A. GUEST ============================
add_heading_custom('A. FITUR UNTUK GUEST (BELUM LOGIN)', level=1)

doc.add_paragraph('Guest adalah pengunjung yang belum melakukan login. Fitur yang tersedia sangat terbatas untuk mendorong pengguna melakukan registrasi.')

add_section_heading('A1. Landing Page / Welcome Page')
doc.add_paragraph(
    'Halaman utama saat pertama kali membuka website. Terdiri dari beberapa section:'
)
add_bullet('Hero Section: Background gradient hijau tua gelap dengan dot pattern overlay. Judul besar "Small Actions, Sustainable Future" dengan aksen italic gold. Tombol CTA "Start Your Journey" dan "Sign In".')
add_bullet('How It Works: 3 cards (Track, Learn, Earn) yang menjelaskan value proposition aplikasi.')
add_bullet('SDG Strip: Daftar semua 17 Sustainable Development Goals dari PBB dengan nomor urut dan nama masing-masing.')
add_bullet('CTA Section: Ajakan mendaftar dengan tombol "Create Your Free Account".')
add_bullet('Footer: Branding dan copyright.')

add_section_heading('A2. Register Akun')
doc.add_paragraph(
    'Form pendaftaran dengan field: nama, email, password, dan konfirmasi password. Setelah berhasil registrasi, user langsung diarahkan ke halaman input body data untuk mengisi berat badan, tinggi, dan kota.'
)

add_section_heading('A3. Login')
doc.add_paragraph(
    'Form login standar dengan email dan password. Mendukung fitur "remember me" dan link lupa password.'
)

add_section_heading('A4. Ganti Bahasa')
doc.add_paragraph(
    'Tombol toggle EN / ID yang tersedia di navbar. Bahasa disimpan dalam session (bukan database) sehingga pengaturan berlaku selama session aktif. Semua teks di website berubah secara dinamis berdasarkan session locale.'
)

doc.add_page_break()

# ============================ B. USER ============================
add_heading_custom('B. FITUR UNTUK USER (SUDAH LOGIN)', level=1)

doc.add_paragraph(
    'Setelah login, user memiliki akses penuh ke semua fitur tracking, edukasi, dan komunitas. Berikut adalah rincian setiap fitur:'
)

# B1 Dashboard
add_section_heading('B1. Dashboard')
doc.add_paragraph(
    'Halaman utama setelah login yang menampilkan ringkasan aktivitas harian user.'
)
add_bullet('Weather Widget: Menampilkan cuaca terkini berdasarkan kota user menggunakan OpenWeatherMap API. Informasi meliputi suhu, kondisi cuaca (cerah/berawan/hujan), dan ikon cuaca.')
add_bullet('Daily Tip: Kutipan atau tips acak tentang kesehatan, nutrisi, atau lingkungan yang muncul setiap kali halaman direfresh.')
add_bullet('5 Stat Cards: Menampilkan total Kalori, Protein, Karbohidrat, Gula, dan Lemak yang sudah dikonsumsi hari ini.')
add_bullet('Progress Bars: Visual progress dari target harian (2000 kcal / 60g protein / 250g carbs / 50g sugar / 65g fat). Setiap bar menunjukkan persentase pencapaian.')
add_bullet('Achievement Badge: Badge level tertinggi yang sudah dicapai user. Dapat diklik untuk menuju halaman achievements detail.')
add_bullet('History Cards: Menampilkan 5 riwayat terakhir. Setiap card menampilkan tanggal dan ringkasan nutrisi/aktivitas. Klik card untuk melihat daily report lengkap.')

# B2 Nutrition
add_section_heading('B2. Nutrisi / Food Tracking')
doc.add_paragraph(
    'Fitur utama untuk mencatat asupan makanan harian. Menyediakan 3 metode input yang berbeda:'
)
add_sub_heading('a. Scan Barcode/Label Makanan')
doc.add_paragraph(
    'User dapat mengupload foto barcode atau label makanan. Sistem akan memproses gambar menggunakan Gemini AI untuk membaca informasi nilai gizi. Hasil scan ditampilkan sebagai preview dan user harus melakukan konfirmasi sebelum data tersimpan. Mendukung 3-tier pencarian otomatis: OpenFoodFacts API, USDA API, dan Gemini sebagai fallback.'
)
add_sub_heading('b. Manual Input')
doc.add_paragraph(
    'Form input manual untuk mencatat makanan. Field: nama makanan, kalori, protein, karbohidrat, gula, lemak. Tipe makanan: snack, makanan_berat, atau minuman. Bisa upload foto makanan (opsional). Sumber tercatat sebagai "manual_input".'
)
add_sub_heading('c. Riwayat Nutrisi')
doc.add_paragraph(
    'Tabel berisi semua log makanan user yang diurutkan berdasarkan tanggal terbaru. Setiap entry menampilkan nama makanan, nilai gizi, tipe makanan, dan waktu. User dapat menghapus entry yang tidak diinginkan.'
)

# B3 Activity
add_section_heading('B3. Aktivitas / Olahraga')
doc.add_paragraph(
    'Fitur untuk mencatat aktivitas fisik dan olahraga.'
)
add_bullet('Log Aktivitas: Form dengan field jenis aktivitas, durasi (menit), jarak (km, opsional), intensitas/pace, kalori terbakar, dan tanggal.')
add_bullet('Riwayat: List semua aktivitas yang pernah dicatat, diurutkan berdasarkan tanggal.')
add_bullet('Hapus: User dapat menghapus entry aktivitas yang tidak diinginkan.')

# B4 Report
add_section_heading('B4. Daily Report')
doc.add_paragraph(
    'Laporan harian yang menggabungkan data nutrisi dan aktivitas. Menampilkan total kalori, protein, karbohidrat, gula, lemak dari makanan, serta total aktivitas (menit dan kalori terbakar). Dilengkapi dengan progress bar visual untuk nutrisi.'
)

# B5 History
add_section_heading('B5. History')
doc.add_paragraph(
    'Menampilkan riwayat harian dalam format card. Setiap card berisi tanggal, total kalori, skor quiz, dan aktivitas. Card dapat diklik untuk menuju halaman report detail pada tanggal tertentu.'
)

# B6 Articles
add_section_heading('B6. Artikel Edukasi')
doc.add_paragraph(
    'Konten edukasi tentang kesehatan dan lingkungan yang dapat dibaca user.'
)
add_bullet('Browse: Grid cards artikel dengan judul, kategori, bahasa, dan excerpt.')
add_bullet('Filter: Berdasarkan kategori (Nutrition & Diet, Disease Prevention, Mental Health, Environmental Health, Fitness & Exercise) dan bahasa (English / Bahasa Indonesia).')
add_bullet('Detail: Halaman artikel dengan konten penuh, kategori badge, bahasa badge, status publish, author, dan source link.')
add_bullet('Komentar: Setiap artikel memiliki kolom komentar. User dapat menulis komentar dan menghapus komentar milik sendiri. Achievement tag dan avatar user ditampilkan di setiap komentar.')

# B7 Discussions
add_section_heading('B7. Forum Diskusi')
doc.add_paragraph(
    'Komunitas diskusi untuk user berinteraksi dan bertukar pikiran.'
)
add_bullet('Index: List semua thread diskusi. Tersedia fitur pencarian dan filter berdasarkan kategori (general, nutrition, sdg, health, tips, lainnya). Thread yang di-pin muncul di bagian atas.')
add_bullet('Create Thread: Membuat diskusi baru dengan field judul, body, dan kategori.')
add_bullet('Show: Halaman detail thread yang menampilkan isi diskusi dan semua reply. Achievement tag dan avatar user ditampilkan di thread dan setiap reply.')
add_bullet('Reply: Membalas thread yang sudah ada. Thread yang di-lock tidak dapat di-reply.')
add_bullet('Delete: User dapat menghapus thread atau reply milik sendiri.')

# B8 Quiz
add_section_heading('B8. Quiz SDG')
doc.add_paragraph(
    'Quiz interaktif seputar Sustainable Development Goals (SDG).'
)
add_bullet('50 Pertanyaan: Tersebar di 5 topics (SDG, Nutrition, Health, Environment, General). Masing-masing topic memiliki 10 soal.')
add_bullet('Filter Topic: Pilihan tab All / SDG / Nutrition / Health / Environment / General untuk memfilter soal.')
add_bullet('Count Selector: Pilihan jumlah soal yang akan dikerjakan: 3, 5, atau 10 soal.')
add_bullet('Progress Bar: Indikator visual yang menunjukkan progres selama mengerjakan quiz.')
add_bullet('Skor Dinamis: Skor per soal dihitung berdasarkan jumlah soal. Contoh: 10 soal = 10 poin per soal, 3 soal = 33.33 poin per soal.')
add_bullet('Review: Setelah quiz selesai, user dapat melihat jawaban mana yang benar (ditandai hijau) dan salah (ditandai merah).')

# B9 Achievements
add_section_heading('B9. Achievements')
doc.add_paragraph(
    'Halaman yang menampilkan semua achievement dan status pencapaian user.'
)
add_bullet('5 Tiers: Eco Starter (level 1), Green Advocate (level 2), Earth Guardian (level 3), SDG Scholar (level 4), Planet Champion (level 5).')
add_bullet('Setiap tier memiliki: logo circle dengan warna khas, nama, deskripsi, dan unlock criteria.')
add_bullet('Status: Ditampilkan "Earned" dengan timestamp jika sudah dicapai, atau "Locked" jika belum.')
add_bullet('Display: Achievement yang sudah di-unlock ditampilkan sebagai colored pill di samping username user di seluruh forum dan komentar. Warna pill berbeda tiap level, mirip Discord role tags.')

# B10 SDG
add_section_heading('B10. SDG Detail')
doc.add_paragraph(
    'Halaman detail untuk masing-masing 17 Sustainable Development Goals. Menampilkan judul, deskripsi lengkap, pentingnya goal tersebut, 3 target spesifik yang dapat diukur, dan 3 aksi nyata yang bisa dilakukan user.'
)

# B11 Profile
add_section_heading('B11. Profil')
doc.add_paragraph(
    'Manajemen akun user secara keseluruhan.'
)
add_bullet('Edit Profil: Mengubah nama dan email.')
add_bullet('Upload Foto Profil: Format jpeg, png, jpg, gif, webp. Maksimal 2MB. Jika tidak ada foto, akan ditampilkan initial circle (huruf pertama dari nama user).')
add_bullet('Remove Foto: Menghapus foto profil dan kembali ke fallback initial circle.')
add_bullet('Body Data: Input berat badan (kg), tinggi (cm), dan kota.')
add_bullet('Detect Location: Mendeteksi lokasi user secara otomatis dari koordinat GPS menggunakan reverse geocoding.')
add_bullet('Ganti Password: Form perubahan password dengan konfirmasi.')
add_bullet('Hapus Akun: Penghapusan akun secara permanen.')

# B12 Body Data
add_section_heading('B12. Body Data')
doc.add_paragraph(
    'Halaman khusus yang muncul setelah registrasi user baru. Form input data tubuh pertama kali: berat badan (kg), tinggi badan (cm), dan kota tempat tinggal.'
)

doc.add_page_break()

# ============================ C. ADMIN ============================
add_heading_custom('C. FITUR ADMIN', level=1)

doc.add_paragraph(
    'Fitur khusus untuk user dengan role admin. Akses melalui route /admin. Admin dapat mengelola seluruh konten dan user di platform.'
)

# C1
add_section_heading('C1. Admin Dashboard')
doc.add_paragraph(
    'Halaman utama admin dengan ringkasan statistik platform. Menampilkan 6 stat cards: Total Users, Total Articles (dengan jumlah published), Today\'s Activities (dengan jumlah active users), Quiz Questions, Comments, dan Discussions. Juga menampilkan Latest Users dan Latest Articles.'
)

# C2
add_section_heading('C2. Manage Articles (CRUD)')
doc.add_paragraph(
    'Pengelolaan penuh untuk artikel edukasi dengan fitur:'
)
add_bullet('Index: Tabel semua artikel dengan filter berdasarkan search (title/excerpt), kategori, dan bahasa.')
add_bullet('Create: Form lengkap dengan field title, slug (auto-generate), kategori, bahasa, excerpt, cover image, content (menggunakan Quill WYSIWYG editor), source URL, author, dan toggle published/draft.')
add_bullet('Edit: Mengubah artikel yang sudah ada.')
add_bullet('Show: Melihat detail artikel.')
add_bullet('Delete: Menghapus artikel.')

# C3
add_section_heading('C3. Manage Users')
doc.add_paragraph(
    'Manajemen user platform:'
)
add_bullet('Index: Tabel semua user dengan fitur search berdasarkan nama atau email. Menampilkan role (Admin/User) dan tanggal join.')
add_bullet('Show: Detail user lengkap meliputi avatar, nama, email, BMI, berat, tinggi, kota, riwayat daily history, dan daftar aktivitas.')

# C4
add_section_heading('C4. Manage Comments')
doc.add_paragraph(
    'Pengelolaan komentar pada artikel:'
)
add_bullet('Index: Tabel semua komentar yang menampilkan isi komentar, nama user, artikel terkait (link), dan tanggal.')
add_bullet('Delete: Menghapus komentar yang tidak sesuai.')

# C5
add_section_heading('C5. Manage Discussions')
doc.add_paragraph(
    'Pengelolaan thread forum diskusi:'
)
add_bullet('Index: Tabel semua thread dengan informasi judul, author, kategori, jumlah reply, status (pinned/locked/active), dan tanggal.')
add_bullet('Pin/Unpin: Mem-pin thread agar muncul di bagian atas daftar.')
add_bullet('Lock/Unlock: Mengunci thread sehingga tidak dapat di-reply.')
add_bullet('View: Melihat thread di halaman frontend (new tab).')
add_bullet('Delete: Menghapus thread beserta semua reply di dalamnya.')

# C6
add_section_heading('C6. Manage Quiz Questions (CRUD)')
doc.add_paragraph(
    'Pengelolaan bank soal quiz:'
)
add_bullet('Index: Tabel semua soal yang menampilkan pertanyaan, opsi jawaban, jawaban benar (dengan badge hijau), dan tanggal dibuat.')
add_bullet('Create/Edit: Form dengan textarea pertanyaan, dynamic options (2-6 opsi dengan tombol tambah/hapus), dan dropdown jawaban benar yang otomatis sync dengan opsi yang diinput.')
add_bullet('Delete: Menghapus soal.')

# C7
add_section_heading('C7. Data Overview')
doc.add_paragraph(
    'Statistik aktivitas pengguna dalam platform:'
)
add_bullet('Period Filter: Pilihan periode Today / This Week / This Month.')
add_bullet('6 Stat Cards: Activities Logged (jumlah aktivitas), Total Minutes (total menit), Calories Burned (total kalori terbakar), Active Users (user aktif), Total Users, dan Avg Quiz Score (rata-rata skor quiz).')
add_bullet('Top Activities: Daftar aktivitas terpopuler diurutkan berdasarkan frekuensi, menampilkan nama aktivitas, total kali dilakukan, dan total menit.')

doc.add_page_break()

# ============================ D. TEKNIS ============================
add_heading_custom('D. FITUR TEKNIS / LINTAS FITUR', level=1)

add_section_heading('D1. Multi Bahasa')
doc.add_paragraph(
    'Mendukung 2 bahasa: English (EN) dan Bahasa Indonesia (ID). Perubahan bahasa dilakukan melalui tombol toggle di navbar dan langsung berlaku tanpa reload halaman penuh. Bahasa disimpan dalam session, bukan database. Semua teks statis menggunakan Laravel localization strings.'
)

add_section_heading('D2. Achievement System')
doc.add_paragraph(
    'Sistem achievement 5 tier yang terinspirasi dari Discord role system. Triggered secara otomatis saat user membuka dashboard atau mengupdate profil. Menggunakan metode syncWithoutDetaching() untuk mencegah duplikasi data pivot. Setiap achievement memiliki warna khas yang ditampilkan sebagai pill di username user di seluruh forum dan komentar. Achievement disimpan di database dengan tabel achievements dan pivot table achievement_user.'
)

add_section_heading('D3. AI Integration')
doc.add_paragraph(
    'Integrasi dengan beberapa layanan AI dan API eksternal:'
)
add_bullet('Gemini AI (Google): Digunakan untuk membaca nilai gizi dari foto barcode/label makanan. Juga sebagai fallback untuk food search ketika OpenFoodFacts dan USDA tidak menemukan hasil.')
add_bullet('OpenWeatherMap: Menyediakan data cuaca untuk weather widget di dashboard.')
add_bullet('Location Service: Reverse geocoding dari koordinat GPS (latitude, longitude) ke nama kota menggunakan WeatherService.')

add_section_heading('D4. Tiga-tier Food Search')
doc.add_paragraph(
    'Sistem pencarian makanan berlapis 3 tingkat untuk memastikan user selalu mendapatkan hasil:'
)
add_bullet('Tingkat 1 - OpenFoodFacts API: Database makanan open-source terbesar. Menggunakan raw PHP curl dengan timeout 8 detik.')
add_bullet('Tingkat 2 - USDA API: Database makanan dari United States Department of Agriculture. Menggunakan GET request dengan api_key query parameter.')
add_bullet('Tingkat 3 - Gemini AI: Estimasi nilai gizi berbasis teks menggunakan AI sebagai langkah terakhir ketika dua API sebelumnya gagal atau timeout.')

# ============================ CLOSING ============================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('--- End of Document ---')
run.font.color.rgb = RGBColor(0x7A, 0x8A, 0x7A)
run.font.size = Pt(10)
run.italic = True

# -- Save --
output_path = os.path.join(os.path.dirname(__file__), 'Daftar_Fitur_EcoLife_Hub.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
